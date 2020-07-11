import docx
import os

replace_dict = {'的': '嘅', '日': '号', '在': '喺', '也': '亦', '他': '佢', '她': '佢', '了': '咗', '是': '系', '没有': '冇',
                '来': '嚟', '着': '住', '看': '睇', '和': '同', '让': '令', '这': '依', '把': '将', '我们': '我哋',
                '他们': '他哋', '她们': '她哋', '它们': '它哋', '还有': '仲有', '里': '入面', '节号': '节日', '这个': '依个',
                '有些': '有D', '找到': '稳到', '回到': '返到', '月月睇': '月月看', '咗解': '了解', '号常': '日常', '这些': '依D',
                '依些': '依D', '生号': '生日', '就像': '就似', '同谐': '和谐', '详同': '详和', '孩子': '小朋友', '还能': '仲可以',}
path = r'E:\我的坚果云\月月看6月配音'
filename_list = os.listdir(path)
print(filename_list)

# 找到合适的文件
new_list = []
for filename in filename_list:
    if filename.endswith('docx'):
        if not filename.split('.')[0].endswith('粤语'):
            new_list.append(filename)

print(new_list)
# 批量替换主程序
for filename in new_list:
    file = docx.Document('{0}\{1}'.format(path, filename))
    print(filename + "的段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

    # 输出每一段的内容
    # for para in file.paragraphs:
    #     print(para.text)

    # 输出段落编号及段落内容
    word_to_txt = ''
    for i in range(len(file.paragraphs)):
        # print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)
        word_to_txt += file.paragraphs[i].text + '\n'
        # 写入文本
        # with open(r'{0}\{1}.txt'.format(path, filename.split('.')[0]), 'w', encoding='utf-8') as f:
        #     f.write(word_to_txt)

    # 提取文本
    # with open(r'{0}\{1}.txt'.format(path, filename.split('.')[0]), 'r', encoding='utf-8') as f:
    #     script = f.read()

    # 批量替换
    for k, v in replace_dict.items():
        word_to_txt = word_to_txt.replace(k, v)

    print('Cantonese:')
    print(word_to_txt)

    # 写入文件
    with open(r'{0}\{1} - 粤语.txt'.format(path, filename.split('.')[0]), 'w', encoding='utf-8') as file:
        file.write(word_to_txt)
