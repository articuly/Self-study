import os
import docx
import tkinter as tk
from tkinter import filedialog
import chardet

cwd = os.getcwd()
replace_dict = {'的': '嘅', '日': '号', '在': '喺', '也': '亦', '他': '佢', '她': '佢', '了': '咗', '是': '系', '没有': '冇',
                '来': '嚟', '着': '住', '看': '睇', '和': '同', '让': '令', '这': '依', '把': '将', '我们': '我哋',
                '他们': '佢哋', '她们': '佢哋', '它们': '佢哋', '还有': '仲有', '里': '入面', '节号': '节日', '这个': '依个',
                '有些': '有D', '找到': '稳到', '回到': '返到', '月月睇': '月月看', '咗解': '了解', '号常': '日常', '这些': '依D',
                '依些': '依D', '生号': '生日', '就像': '就似', '同谐': '和谐', '详同': '详和', '孩子': '小朋友', '还能': '仲可以', }


def divide_filenames(filenames_paths):
    docx_list = []
    txt_list = []
    for filename in filenames_paths:
        if filename.endswith('docx'):
            if not filename.split('.')[-2].endswith('粤语'):
                docx_list.append(filename)
        else:
            if not filename.split('.')[-2].endswith('粤语'):
                txt_list.append(filename)
    return docx_list, txt_list


def docx2cantonese(docx_list):
    for file_path in docx_list:
        try:
            file = docx.Document(file_path)
        except Exception as e:
            print(e)
            print('提示：Docx文档里可能没有内容')
        else:
            # 转成分行的字符串
            cantonese_txt = ''
            for line in range(len(file.paragraphs)):
                cantonese_txt += file.paragraphs[line].text + '\n'
            # print(cantonese_txt)
            # 批量替换
            for k, v in replace_dict.items():
                cantonese_txt = cantonese_txt.replace(k, v)
            # print(cantonese_txt)
            # 写入文件
            global cwd
            filepath, full_filename = os.path.split(file_path)
            filename, ext = os.path.splitext(full_filename)
            with open(r'{0}\{1} - 粤语.txt'.format(cwd, filename), 'w', encoding='utf-8') as file:
                file.write(cantonese_txt)
            print(r'{0}\{1} - 粤语.txt 保存成功'.format(cwd, filename))


def txt2cantonese(txt_list):
    for file_path in txt_list:
        with open(file_path, 'rb') as f:
            cur_encoding = chardet.detect(f.read())['encoding']
            # print(cur_encoding)  # 当前文件编码
        # 用获取的编码读取该文件而不是python3默认的utf-8读取。
        try:
            with open(file_path, encoding=cur_encoding) as file:
                cantonese_txt = file.read()
        except Exception as e:
            print(e)
        else:
            # 批量替换
            for k, v in replace_dict.items():
                cantonese_txt = cantonese_txt.replace(k, v)
            # 写入文件
            global cwd
            filepath, full_filename = os.path.split(file_path)
            filename, ext = os.path.splitext(full_filename)
            with open(r'{0}\{1} - 粤语.txt'.format(cwd, filename), 'w', encoding='utf-8') as file:
                file.write(cantonese_txt)
            print(r'{0}\{1} - 粤语.txt 保存成功'.format(cwd, filename))


if __name__ == '__main__':
    root = tk.Tk()  # 创建一个Tkinter.Tk()实例
    root.withdraw()  # 将Tkinter.Tk()实例隐藏
    filenames_paths = filedialog.askopenfilenames(title='请选择要转换成粤语的文档', initialdir='%USERPROFILE%',
                                                  filetypes=[("支持文档", ".txt .docx")],
                                                  defaultextension='.txt', multiple=True)
    print(filenames_paths)
    docx_list, txt_list = divide_filenames(filenames_paths)
    # print(docx_list, txt_list)
    docx2cantonese(docx_list)
    txt2cantonese(txt_list)
